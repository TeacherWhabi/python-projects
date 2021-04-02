# Import Document to manipulate doc files
from docx import Document
# Import Inches to resize the pictures
from docx.shared import Inches
# Import Text to Speech Library
import pyttsx3

# Adjust TTS Properties
converter = pyttsx3.init()
converter.setProperty('rate', 100)
converter.setProperty('volume', 0.9)


# TTS function
def speak(text):
    pyttsx3.speak(text)

# Start working on the document
document = Document()

# Insert a picture into the document
document.add_picture(
    'Profile.jpeg', 
    width=Inches(2.0)
)

# Input info from user
name = input('What is your name? ')
speak('Hello '+ name + ' how are you today?')

speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

# Write into the document
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me infos + Heading Title
document.add_heading('About Me')
about_me = input('Tell me about yourself: ')
document.add_paragraph(about_me)

# Work Experience + Heading Title
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('From Date: ')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '_' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# More Experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        
        p = document.add_paragraph()

        company = input('Enter company: ')
        from_date = input('From Date: ')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '_' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ': ')
        p.add_run(experience_details)

    else:
        break
    
# Skills + Heading Title
document.add_heading('Skills')

skill = input('Enter skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    add_skills = input('Do you want to add another skill? Yes or No: ')
    if add_skills.lower() == 'yes':
        skill = input('Enter skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'  
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Python"
        
# Save cv.docx
document.save('cv.docx')



